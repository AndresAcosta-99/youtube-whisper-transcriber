"""
Script de Transcripción Automática de Videos de YouTube
========================================================

Descarga videos de YouTube, extrae el audio, y genera transcripciones
en español usando el modelo Whisper de OpenAI.

Uso:
    - Con CSV: python3 script.py videos.csv
    - Manual:  python3 script.py

Requisitos:
    - yt-dlp, whisper, torch, pandas, python-docx, ffmpeg
"""

import yt_dlp
import subprocess
import whisper
import sys
import pandas as pd
import os
import torch
from docx import Document
import re


def main():
    """
    Función principal que coordina el proceso de transcripción.
    """
    if len(sys.argv) == 2:
        videos = extract_url(sys.argv[1])
    else:
        videos = [input("Enter the video URL: ")]
    
    log_file = "transcriptions_log.csv"

    device = get_device()
    
    print("\n📦 Cargando modelo Whisper 'medium'...")
    model = whisper.load_model("medium", device=device)
    print("✓ Modelo cargado\n")

    # Contadores para resumen final
    total_videos = len(videos)
    exitosos = 0
    fallidos = 0
    
    for idx, url in enumerate(videos, 1):
        # Inicializar variables para limpieza segura
        input_file = None
        output_file = None
        
        # Indicador de progreso
        print("=" * 70)
        print(f"📹 VIDEO {idx}/{total_videos}")
        print(f"🔗 {url}")
        print("=" * 70)
        
        try:
            # Descargar
            print(f"\n[{idx}/{total_videos}] Paso 1/4: Descargando audio...")
            info = download_video(url)
            print(f"✓ Descargado: {info['title'][:50]}...")
            
            # Convertir
            print(f"[{idx}/{total_videos}] Paso 2/4: Convirtiendo a WAV...")
            input_file, output_file = convert_to_wav(info)
            print(f"✓ Convertido")
            
            # Transcribir
            print(f"[{idx}/{total_videos}] Paso 3/4: Transcribiendo...")
            result = model.transcribe(output_file, language="es", fp16=False)
            print(f"✓ Transcripción completada ({len(result['text'])} caracteres)")
            
            # Guardar
            print(f"[{idx}/{total_videos}] Paso 4/4: Guardando documento...")
            safe_title = re.sub(r'[<>:"/\\|?*]', '_', info['title'])
            word(result, safe_title + ".docx")
            
            # Registrar éxito
            registro(info, url, log_file)
            exitosos += 1
            
            print(f"\n✅ VIDEO {idx}/{total_videos} COMPLETADO\n")
            
        except Exception as e:
            print(f"\n❌ ERROR en video {idx}/{total_videos}: {str(e)}\n")
            registro(e, url, log_file)
            fallidos += 1
            
        finally:
            # Limpieza garantizada
            clean(input_file, output_file)
    
    # Resumen final
    print("\n" + "=" * 70)
    print("🏁 PROCESO COMPLETADO")
    print("=" * 70)
    print(f"✅ Exitosos: {exitosos}/{total_videos}")
    print(f"❌ Fallidos:  {fallidos}/{total_videos}")
    print(f"📊 Log guardado en: {log_file}")
    print("=" * 70 + "\n")


def extract_url(csv):
    """Extrae URLs desde un archivo CSV"""
    if not csv.endswith('.csv'):
        raise ValueError("El archivo proporcionado no es un archivo CSV válido.")
    if not os.path.exists(csv):
        raise FileNotFoundError(f"No se encontró el archivo: {csv}")
    
    df = pd.read_csv(csv) 
    
    if 'url' not in df.columns:
        raise ValueError("El archivo CSV no contiene una columna 'url'.")
    
    urls = []
    for link in df['url'].tolist():
        if pd.notna(link) and str(link).strip():
            urls.append(link.strip())
    
    if not urls:
        raise ValueError("No se encontraron URLs válidas en el CSV.")
    
    print(f"📋 Se encontraron {len(urls)} video(s) para procesar")
    return urls


def get_device():
    """Detecta automáticamente si hay GPU disponible"""
    if torch.cuda.is_available():
        device = "cuda"
        print(f"✓ GPU detectada: {torch.cuda.get_device_name(0)}")
        print(f"  Memoria GPU disponible: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.1f} GB")
    else:
        device = "cpu"
        print("⚠ No se detectó GPU, usando CPU")
    return device


def download_video(url):
    """Descarga el audio de un video de YouTube"""
    ydl_opts = {
        'format': 'bestaudio',        
        'outtmpl': '%(title)s.%(ext)s',
        'quiet': False,                
        'noplaylist': True             
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
    return info


def convert_to_wav(info):
    """Convierte el audio descargado a formato WAV optimizado para Whisper"""
    input_file = info['requested_downloads'][0]['filepath']
    
    # Sanitizar título para evitar errores con caracteres especiales
    safe_title = re.sub(r'[<>:"/\\|?*]', '_', info['title'])
    output_file = safe_title + ".wav"

    comando = [
        "ffmpeg",
        "-y",  # Sobrescribir sin preguntar
        "-i", input_file,
        "-ac", "1",
        "-ar", "16000",
        "-c:a", "pcm_s16le",
        "-af", "loudnorm",
        output_file,
    ]

    subprocess.run(comando, check=True, capture_output=True)
    return input_file, output_file


def word(result, documento="transcripcion-med.docx"):
    """
    Guarda el texto transcrito en un documento Word con timestamps.
    
    Args:
        result (dict): Resultado completo de Whisper con 'text' y 'segments'
        documento (str): Nombre del archivo .docx
    """
    doc = Document()
    
    # Título del documento
    doc.add_heading('Transcripción de Audio', 0)
    
    # Agregar transcripción con timestamps si están disponibles
    if 'segments' in result:
        for segment in result['segments']:
            # Formatear timestamp
            start_time = format_timestamp(segment['start'])
            end_time = format_timestamp(segment['end'])
            
            # Agregar párrafo con timestamp en negrita y texto normal
            p = doc.add_paragraph()
            p.add_run(f"[{start_time} → {end_time}] ").bold = True
            p.add_run(segment['text'].strip())
    else:
        # Fallback: si no hay segments, usar texto completo
        doc.add_paragraph(result['text'])
    
    doc.save(documento)
    print(f"✓ Guardado en: {documento}")


def format_timestamp(seconds):
    """
    Convierte segundos a formato timestamp legible.
    
    Args:
        seconds (float): Tiempo en segundos
    
    Returns:
        str: Timestamp en formato "HH:MM:SS" o "MM:SS"
    """
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    else:
        return f"{minutes:02d}:{secs:02d}"


def clean(input_file, output_file):
    """Elimina archivos temporales de forma segura"""
    archivos_eliminados = []
    
    for archivo in [input_file, output_file]:
        if archivo and os.path.exists(archivo):
            try:
                os.remove(archivo)
                archivos_eliminados.append(os.path.basename(archivo))
            except OSError as e:
                print(f"⚠ No se pudo eliminar {archivo}: {e}")
    
    if archivos_eliminados:
        print(f"🗑️  Eliminados: {', '.join(archivos_eliminados)}")


def registro(info, url, output_csv="transcriptions_log.csv"):
    """Registra el resultado del procesamiento en un CSV"""
    if isinstance(info, Exception):
        row = {
            "title": "N/A",
            "upload_date": "N/A",
            "duration": "N/A",
            "view_count": "N/A",
            "channel": "N/A",
            "URL": url,
            "status": f"Error: {str(info)}"
        }
    else:
        # Formatear fecha: de "20081010" a "2008-10-10"
        upload_date = info.get('upload_date', 'N/A')
        if upload_date != 'N/A' and len(upload_date) == 8:
            upload_date = f"{upload_date[0:4]}-{upload_date[4:6]}-{upload_date[6:8]}"
        
        # Formatear duración: de segundos a "HH:MM:SS"
        duration = info.get('duration', 'N/A')
        if isinstance(duration, (int, float)):
            hours = int(duration // 3600)
            minutes = int((duration % 3600) // 60)
            seconds = int(duration % 60)
            duration = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        
        row = {
            "title": info.get('title', 'N/A'),
            "upload_date": upload_date,
            "duration": duration,
            "view_count": info.get('view_count', 'N/A'),
            "channel": info.get('uploader', 'N/A'),
            "URL": url,
            "status": "Success",
        }
    
    df = pd.DataFrame([row])
    write_header = not os.path.exists(output_csv)

    df.to_csv(
        output_csv,
        mode="a",
        header=write_header,
        index=False
    )


if __name__ == "__main__":
    main()
